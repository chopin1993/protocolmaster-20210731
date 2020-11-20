#encoding:utf-8
import docx
import logging
import time
import os
__version__ = "V1.0"


def get_config_file(name):
    return os.path.join(os.path.dirname(__file__), "../resource", name)

class DocxEngine(object):
    def __init__(self, name=None):
        self.name = None
        self.document = docx.Document(get_config_file("template_test.docx"))
        styles = [s for s in self.document.styles]
        if name is not None:
            self.write_doc_head(name)

    def write_doc_head(self, name):
        self.document = docx.Document(get_config_file("template_test.docx"))
        self.name = name
        self.document.add_heading("\n\n\n", 0)
        self.document.add_heading("自动测试", 0)
        self.document.add_heading(name, 0)
        self.document.add_heading("\n\n\n\n\n", 0)
        time_str = "测试时间:{0}".format(time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time())))
        self.document.add_paragraph(" "*35 + time_str)
        engine_version = "测试引擎版本:{0}".format(__version__)
        self.document.add_paragraph(" " * 35 + engine_version)
        self.document.add_page_break()

    def write_summary(self, total, passed, fails, all_infos):
        fails_cnt = len(fails)
        self.document.add_heading("概要", 1)
        txt = "总测试用例:{0} 通过:{1} 失败:{2}".format(total, passed, fails_cnt)
        self.document.add_paragraph(txt, style="important")
        self.document.add_heading("失败用例汇总", 2)
        table = self.document.add_table(rows=1, cols=2,style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '失败用例'
        hdr_cells[1].text = '失败原因'
        if fails:
            for case in fails:
                case.write_fail_table(table)
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "无"
            row_cells[1].text = "无"
        self.document.add_heading("所有测试用例汇总", 2)
        table = self.document.add_table(rows=1, cols=3,style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '测试组'
        hdr_cells[1].text = '测试用例'
        hdr_cells[2].text = '状态'
        if all_infos:
            for info in all_infos:
                info.write_summary_table(table)
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "无"
            row_cells[1].text = "无"
            row_cells[2].text = "无"
        self.document.add_page_break()

    def write_detail(self, infos):
        self.document.add_heading("测试详情", 1)
        for case in infos:
            case.write_doc(self, True)

    def start_group(self, name, brief=None):
        self.document.add_heading(name, 2)
        if brief is not None:
            self.add_tag_msg("", "doc", brief)

    def end_group(self, name):
        pass

    def start_test(self, name, brief=None):
        self.document.add_heading(name, 3)
        if brief is not None:
            self.add_tag_msg("", "doc", brief)

    def end_test(self, name):
        pass

    def add_tag_msg(self, name, tag, msg, timestr=""):
        style = "code"
        if tag == "snd":
            body = "{0} {1}::{2} -> {3}".format(timestr, name, tag, msg)
            style = "code"
        elif tag in ["rcv","rev"]:
            body = "{0} {1}::{2} <- {3}".format(timestr, name, tag, msg)
            style = "rcv"
        elif tag == "annotation":
            body = msg
            style = "annotation"
        elif tag in ["expect success","sucess","success"]:
            body = msg
            style = "success"
        elif tag in ["expect fail","fail","exception"]:
            body = msg
            style = "fail"
        elif tag in ['doc']:
            body = msg
            style = "doc"
        else:
            logging.info("no handle tag %s",tag)
            body = "{0}::{1} {2}".format(name, tag, msg)
            style = "annotation"
        self.document.add_paragraph(body, style=style)

    def add_fail(self, body):
        self.document.add_paragraph(body, style='fail')

    def save_doc(self,output_dir):
       output_dir = os.path.join(output_dir, "测试报告")
       if not os.path.exists(output_dir):
           os.mkdir(output_dir)
       time_str = time.strftime('%Y-%m-%d-%H-%M-%S', time.localtime(time.time()))
       name = self.name+"_"+time_str+".docx"
       name = os.path.join(output_dir,name)
       logging.info("generate doc %s",name)
       self.document.save(name)


if __name__ == "__main__":
    doc = DocxEngine("ESSN-OIP-19900101(V1.2)")
    doc.start_group("基本报文管理")
    doc.start_test("版本测试")
    doc.add_normal("send 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("analyse 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("rcv 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("analyse 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.end_test("版本测试")
    doc.end_group("基本报文管理")

    doc.start_group("上报测试")
    doc.start_test("上电上报")
    doc.add_normal("send 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("analyse 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("rcv 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.add_normal("analyse 73 33 33 33 33 33 33 33 33 33 3 ")
    doc.end_test("上电上报")
    doc.end_group("上报测试")
    doc.save_doc()


